# # #Python demo1.py

# from flask import Flask, request, send_file, jsonify
# from langchain_groq import ChatGroq
# from langchain_core.output_parsers import StrOutputParser
# from langchain_core.prompts import ChatPromptTemplate
# from docx import Document
# from flask_cors import CORS
# import pandas as pd
# import os
# from dotenv import load_dotenv
# load_dotenv()  # Load environment variables from .env file
# app = Flask(__name__)
# CORS(app)  # Enable CORS for frontend communication

# llm = ChatGroq(temperature=0.5, groq_api_key=os.getenv("GROQ_API_KEY"), model_name="llama3-8b-8192")
# def generate_vsq_questions(subject_name, syllabus, num_vsq):
#     prompt_template = """
#     You are an expert in creating educational content. Based on the following inputs, generate {num_vsq} Very Short Question worth 2 Marks for an examination.
#     The questions should be using these question starter 
#     Define...

#     What is...

#     List...

#     State...

#     Name...

#     Identify...

#     Mention...

#     Give examples of...

#     and the 
#     Subject: {subject_name}
#     Syllabus: {syllabus}

#     Please write very short answer for the following questions.
#     Please do not provide the answers to the questions.

#     Q1. _________________
    
#     Q2. _________________
#     ...
#     Generate the questions in the  same format.
#     """

#     prompt = prompt_template.format(
#         subject_name=subject_name,
#         syllabus=syllabus,
#         num_vsq=num_vsq,
#     )

#     chain = (
#         ChatPromptTemplate.from_template(prompt)
#         | llm
#         | StrOutputParser()
#     )

#     return chain.invoke({})

# def generate_short_questions(subject_name, syllabus, num_short):
#     prompt_template = """
#     You are an expert in creating educational content. Based on the following inputs, generate {num_short} short answer questions for an examination.The Question Should be of Worth of 4 Marks
#     The questions should be using these question starter 
#     Explain...

#     Describe...

#     Differentiate between...

#     Illustrate with examples...

#     Compare...

#     Write short notes on...

#     How does... work?

#     Discuss...

#     and the 
#     Subject: {subject_name}
#     Syllabus: {syllabus}

#     Please do not provide the answers to the questions.
#     Do not provide any notes to the questions.

#     Q1. _________________
    
#     Q2. _________________
    
#     ...
    
#     Generate the questions in the same format.
#     """

#     prompt = prompt_template.format(
#         subject_name=subject_name,
#         syllabus=syllabus,
#         num_short=num_short
#     )

#     chain = (
#         ChatPromptTemplate.from_template(prompt)
#         | llm
#         | StrOutputParser()
#     )

#     return chain.invoke({})

# def generate_long_questions(subject_name, syllabus, num_long):
#     prompt_template = """
#     You are an expert in creating educational content. Based on the following inputs, generate {num_long} long answer questions for an examination. The Question SHould be worth 12 Marks.
#     The questions should be using these question starter 
#     Analyze...

#     Evaluate...

#     Justify...

#     Design a...

#     Develop a...

#     Write a detailed note on...

#     Compare and evaluate...

#     Propose a solution for...
#     and the
#     Subject: {subject_name}
#     Syllabus: {syllabus}
#     Please do not provide the answers to the questions.
#     Do not provide any notes to the questions.

#     Q1. _________________
    
#     Q2. _________________
    
#     ...
    
#     Generate the questions in the same format.
#     """

#     prompt = prompt_template.format(
#         subject_name=subject_name,
#         syllabus=syllabus,
#         num_long=num_long
#     )

#     chain = (
#         ChatPromptTemplate.from_template(prompt)
#         | llm
#         | StrOutputParser()
#     )

#     return chain.invoke({})

# def generate_case_questions(subject_name, syllabus, num_case):
#     prompt_template = """
#     You are an expert in creating educational content. Based on the following inputs, generate {num_case} case study questions for an examination.

#     Subject: {subject_name}
#     Syllabus: {syllabus}

#     Generate a long full case study paragraph combining multiple topics from syllabus and ask questions accordingly questions should be of 12 marks combined. Case Study Paragraph Should be 200-250 words atleast. 
#     Do not provide answers to the questions.

#     case1__________________
#     ________________________
#     ________________________
    
#     Q1. _________________
#     Q2. _________________

#     case2__________________
#     _______________________
#     _______________________

#     Q1. _________________
#     Q2. _________________
#     ...
    
#     Generate the questions in the same format.
#     """

#     prompt = prompt_template.format(
#         subject_name=subject_name,
#         syllabus=syllabus,
#         num_case=num_case
#     )

#     chain = (
#         ChatPromptTemplate.from_template(prompt)
#         | llm
#         | StrOutputParser()
#     )

#     return chain.invoke({})
# SetsNum=0
# @app.route('/', methods=['GET'])
# def home():
#     return "API is running"
# @app.route('/generate', methods=['POST'])
# def generate():
#     try:
#         # Get form data
#         subject = request.form.get("subject", "")
#         vsq_num = int(request.form.get("vsq_num", 0))
#         sq_num = int(request.form.get("sq_num", 0))
#         long_num = int(request.form.get("long_num", 0))
#         case_num = int(request.form.get("case_num", 0))
#         setsNum = int(request.form.get("setsNum", 1))  # Default to 1 set

#         print(f"setsNum: {setsNum}")  # Debugging

#         # Read uploaded syllabus file
#         file = request.files["syllabus"]
#         df = pd.read_excel(file)

#         # Validate file structure
#         if 'Marks' not in df.columns or 'Topics' not in df.columns:
#             return jsonify({"error": "Invalid file format. 'Marks' or 'Topics' column is missing."}), 400

#         # Filter topics based on marks
#         filter2 = df[df['Marks'] == '2 Marks']
#         filter4 = df[df['Marks'] == '4 Marks']
#         filter12 = df[df['Marks'] == '12 Marks']
#         filtercase = df[df['Marks'] == '12 Marks Case Study']

#         if filter2.empty or filter4.empty or filter12.empty or filtercase.empty:
#             return jsonify({"error": "No data found in the uploaded file for the specified marks."}), 400

#         vsqtopics = filter2['Topics']
#         sqtopics = filter4['Topics']
#         longtopics = filter12['Topics']
#         casetopics = filtercase['Topics']

#         # Debugging
#         print(f"VSQ Topics: {vsqtopics}")
#         print(f"Short Topics: {sqtopics}")
#         print(f"Long Topics: {longtopics}")
#         print(f"Case Topics: {casetopics}")

#         # Create a directory to store the generated files
#         output_dir = "generated_files"
#         os.makedirs(output_dir, exist_ok=True)
#         output_files = []
#           # Debugging
#         doc = Document()
#         doc.add_heading(f"Generated Questions for {subject}")
#         # Generate documents
#         for i in range(1, setsNum + 1):
           
#             doc.add_paragraph(f"Set {i}")
#             if vsq_num > 0:
#                 doc.add_paragraph("Very Short Questions:")
#                 doc.add_paragraph(generate_vsq_questions(subject, vsqtopics, vsq_num))

#             if sq_num > 0:
#                 doc.add_paragraph("Short Questions:")
#                 doc.add_paragraph(generate_short_questions(subject, sqtopics, sq_num))

#             if long_num > 0:
#                 doc.add_paragraph("Long Questions:")
#                 doc.add_paragraph(generate_long_questions(subject, longtopics, long_num))

#             if case_num > 0:
#                 doc.add_paragraph("Case Study Questions:")
#                 doc.add_paragraph(generate_case_questions(subject, casetopics, case_num))
#             doc.add_page_break()

#         filename = f"output.docx"
#         doc.save(filename)
#         return send_file(filename, as_attachment=True)
#     except Exception as e:
#         print(f"Error: {str(e)}")  # Debugging
#         return jsonify({"error": f"Error processing request: {str(e)}"}), 400
# if __name__ == '__main__':
#     port = int(os.environ.get('PORT', 5000))
#     app.run(host='0.0.0.0', port=port)
from flask import Flask, request, send_file, jsonify
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from docx import Document
from flask_cors import CORS
import pandas as pd
import os
from dotenv import load_dotenv

load_dotenv()  # Load environment variables from .env file (for local dev)

app = Flask(__name__)
CORS(app)  # Enable CORS

llm = ChatGroq(
    temperature=0.5,
    groq_api_key=os.getenv("GROQ_API_KEY"),
    model_name="llama3-8b-8192"
)

def generate_vsq_questions(subject_name, syllabus, num_vsq):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_vsq} Very Short Question worth 2 Marks for an examination.
    The questions should be using these question starter 
    Define...

    What is...

    List...

    State...

    Name...

    Identify...

    Mention...

    Give examples of...

    and the 
    Subject: {subject_name}
    Syllabus: {syllabus}

    Please write very short answer for the following questions.
    Please do not provide the answers to the questions.

    Q1. _________________

    Q2. _________________
    ...
    Generate the questions in the same format.
    """
    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_vsq=num_vsq,
    )
    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )
    return chain.invoke({})

def generate_short_questions(subject_name, syllabus, num_short):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_short} short answer questions for an examination.The Question Should be of Worth of 4 Marks
    The questions should be using these question starter 
    Explain...

    Describe...

    Differentiate between...

    Illustrate with examples...

    Compare...

    Write short notes on...

    How does... work?

    Discuss...

    and the 
    Subject: {subject_name}
    Syllabus: {syllabus}

    Please do not provide the answers to the questions.
    Do not provide any notes to the questions.

    Q1. _________________

    Q2. _________________

    ...

    Generate the questions in the same format.
    """
    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_short=num_short
    )
    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )
    return chain.invoke({})

def generate_long_questions(subject_name, syllabus, num_long):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_long} long answer questions for an examination. The Question Should be worth 12 Marks.
    The questions should be using these question starter 
    Analyze...

    Evaluate...

    Justify...

    Design a...

    Develop a...

    Write a detailed note on...

    Compare and evaluate...

    Propose a solution for...
    and the
    Subject: {subject_name}
    Syllabus: {syllabus}
    Please do not provide the answers to the questions.
    Do not provide any notes to the questions.

    Q1. _________________

    Q2. _________________

    ...

    Generate the questions in the same format.
    """
    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_long=num_long
    )
    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )
    return chain.invoke({})

def generate_case_questions(subject_name, syllabus, num_case):
    prompt_template = """
    You are an expert in creating educational content. Based on the following inputs, generate {num_case} case study questions for an examination.

    Subject: {subject_name}
    Syllabus: {syllabus}

    Generate a long full case study paragraph combining multiple topics from syllabus and ask questions accordingly. Questions should be of 12 marks combined. Case Study Paragraph Should be 200-250 words at least. 
    Do not provide answers to the questions.

    case1__________________
    ________________________
    ________________________

    Q1. _________________
    Q2. _________________

    case2__________________
    _______________________
    _______________________

    Q1. _________________
    Q2. _________________
    ...

    Generate the questions in the same format.
    """
    prompt = prompt_template.format(
        subject_name=subject_name,
        syllabus=syllabus,
        num_case=num_case
    )
    chain = (
        ChatPromptTemplate.from_template(prompt)
        | llm
        | StrOutputParser()
    )
    return chain.invoke({})

@app.route('/', methods=['GET'])
def home():
    return "API is running"

@app.route('/generate', methods=['POST'])
def generate():
    try:
        print("Request form:", request.form)
        print("Request files:", request.files)

        # Get file safely
        file = request.files.get("syllabus")
        if not file:
            print("No syllabus file uploaded!")
            return jsonify({"error": "Syllabus file is required"}), 400

        try:
            df = pd.read_excel(file)
            print("Excel columns:", df.columns)
        except Exception as e:
            print("Error reading Excel file:", e)
            return jsonify({"error": "Failed to read syllabus file"}), 400

        if 'Marks' not in df.columns or 'Topics' not in df.columns:
            return jsonify({"error": "Invalid file format. 'Marks' or 'Topics' column is missing."}), 400

        # Get form values
        subject = request.form.get("subject", "")
        vsq_num = int(request.form.get("vsq_num", 0))
        sq_num = int(request.form.get("sq_num", 0))
        long_num = int(request.form.get("long_num", 0))
        case_num = int(request.form.get("case_num", 0))
        setsNum = int(request.form.get("setsNum", 1))
        print(f"setsNum: {setsNum}")

        # Filter topics by marks
        filter2 = df[df['Marks'] == '2 Marks']
        filter4 = df[df['Marks'] == '4 Marks']
        filter12 = df[df['Marks'] == '12 Marks']
        filtercase = df[df['Marks'] == '12 Marks Case Study']

        if filter2.empty or filter4.empty or filter12.empty or filtercase.empty:
            return jsonify({"error": "No data found in the uploaded file for the specified marks."}), 400

        vsqtopics = filter2['Topics'].tolist()
        sqtopics = filter4['Topics'].tolist()
        longtopics = filter12['Topics'].tolist()
        casetopics = filtercase['Topics'].tolist()

        print(f"VSQ Topics: {vsqtopics}")
        print(f"Short Topics: {sqtopics}")
        print(f"Long Topics: {longtopics}")
        print(f"Case Topics: {casetopics}")

        output_dir = "generated_files"
        os.makedirs(output_dir, exist_ok=True)

        doc = Document()
        doc.add_heading(f"Generated Questions for {subject}")

        for i in range(1, setsNum + 1):
            doc.add_paragraph(f"Set {i}")
            if vsq_num > 0:
                doc.add_paragraph("Very Short Questions:")
                doc.add_paragraph(generate_vsq_questions(subject, vsqtopics, vsq_num))

            if sq_num > 0:
                doc.add_paragraph("Short Questions:")
                doc.add_paragraph(generate_short_questions(subject, sqtopics, sq_num))

            if long_num > 0:
                doc.add_paragraph("Long Questions:")
                doc.add_paragraph(generate_long_questions(subject, longtopics, long_num))

            if case_num > 0:
                doc.add_paragraph("Case Study Questions:")
                doc.add_paragraph(generate_case_questions(subject, casetopics, case_num))

            doc.add_page_break()

        filename = os.path.join(output_dir, "output.docx")
        doc.save(filename)

        return send_file(filename, as_attachment=True)

    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({"error": f"Error processing request: {str(e)}"}), 400

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=True)
